import os
import sys
from decimal import Decimal
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement
import matplotlib.pyplot as plt
from tqdm import tqdm

plt.rcParams["font.family"] = "Times New Roman"

df = pd.read_csv("regression_results_2.csv")
df_cutoff = pd.read_csv("cutoff_results_2.csv")
word_document = Document()
out_fname = "regression_results_2.docx"

style = word_document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

reg_model_lst = ["logist"]#["linear", "logist"]
digits = 2
ci_coef = 1.96

output_name_dct_raw = {
    "self_rated_1": {"Name": "Self Rated 1", "Category": "Self Rated Health", "Model": "logist"},
    "self_rated_2": {"Name": "Self Rated 2", "Category": "Self Rated Health", "Model": "logist"},
}

output_name_dct = {}
for model in reg_model_lst:
    output_name_dct[model] = {}
    for key in output_name_dct_raw:
        if output_name_dct_raw[key]["Model"] == model:
            cat = output_name_dct_raw[key]["Category"]
            name = output_name_dct_raw[key]["Name"]
            if cat not in output_name_dct[model]:
                output_name_dct[model][cat] = []
            output_name_dct[model][cat].append((key, name))

def tidy_pval(pval):
    pval = float(pval)
    if pval < 1e-3:
        return "p-trend<0.001"
    else:
        return "p-trend={:.3f}".format(pval)

def tidy_pval2(pval):
    pval = float(pval)
    if pval < 1e-3:
        return "<0.001"
    else:
        return "{:.3f}".format(pval)
    #return f"p = {Decimal(str(pval)):.2E}"

def set_cell_background(cell, fill):
    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd') # add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)
    cell_properties.append(cell_shading)

def add_table(word_document, rows_lst, headline, tableId, yname_lst = None, subhead = None):
    p = word_document.add_paragraph()
    runner = p.add_run(headline)
    runner.bold = True
    if subhead is not None:
            p = word_document.add_paragraph()
            runner = p.add_run(subhead)
            runner.bold = True
    table = word_document.add_table(rows = 0, cols = len(rows_lst[0][0]))
    header_lst = []
    row_idx = 0
    if yname_lst is None:
        yname_lst = [None] * len(rows_lst)
    for rows, yname in zip(rows_lst, yname_lst):
        if yname is not None:
            new_row = table.add_row().cells
            table.cell(row_idx, 1).merge(table.cell(row_idx, len(rows[0]) - 1))
            table.rows[row_idx].cells[1].text = yname
            table.rows[row_idx].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_lst.append(row_idx)
            row_idx += 1
        header_lst.append(row_idx)
        for row in rows:
            new_row = table.add_row().cells
            for i in range(len(row)):
                txt = str(row[i])
                if "p-trend" in txt and len(txt.split("\n")) > 1:
                    p_trend = txt.split("\n")[-1]
                    txt = "\n".join(txt.split("\n")[:-1]).strip()
                    paragraphs = new_row[i].paragraphs
                    p = new_row[i].paragraphs[0]
                    runner = p.add_run(txt)
                    p = new_row[i].add_paragraph()
                    runner = p.add_run(p_trend)
                    runner.font.italic = True
                else:
                    p_trend = None
                    new_row[i].text = str(row[i])
            row_idx += 1
    table.autofit = True
    table.style = "TableGrid"
    for j in header_lst:
        for i in range(len(table.rows[j].cells)):
            set_cell_background(table.rows[j].cells[i], "E5E8E8")

def varname_map(input):
    dct = {
        "IV60": "Intra-daily Variability (IV60)",
        "L5_cnt": "L5 Counts",
        "RA": "Relative Amplitude",
        "L5_midpoint": "L5 Midpoint",
        "M10_midpoint": "M10 Midpoint",
        "IS60": "Inter-daily Stability (IS60)",
        "M10_cnt_100": "M10 Counts (x100)",
        "M10_cnt": "M10 Counts",
        "RIDAGEYR": "Age",
        "alcohol": "Alcohol (# Drinks/Day)",
    }
    if input in dct:
        return dct[input]
    else:
        return input.replace("_", " ").title()

def varname_map2(input):
    dct = {
        "IV60": "Intra-daily\nVariability",
        "L5_cnt": "L5 Counts",
        "RA": "Relative\nAmplitude",
        "L5_midpoint": "L5 Midpoint",
        "M10_midpoint": "M10 Midpoint",
        "IS60": "Inter-daily\nStability",
        "M10_cnt_100": "M10 Counts",
        "M10_cnt": "M10 Counts",
        "RIDAGEYR": "Age",
        "alcohol": "Alcohol (# Drinks/Day)",
    }
    if input in dct:
        return dct[input]
    else:
        return input.replace("_", " ").title()

def compute_bxp_stat(odds_ratio, std):
    med = np.exp(odds_ratio)
    q1 = np.exp(odds_ratio - std)
    q3 = np.exp(odds_ratio + std)
    min = np.exp(odds_ratio - ci_coef * std)
    max = np.exp(odds_ratio + ci_coef * std)
    dct = {"med": med, "q1": q1, "q3": q3, "whislo": min, "whishi": max}
    return dct

def plot_confidence_interval(x, stats_dct, color='#2187bb', horizontal_line_width=0.25, color_box = "#77736c", manual_adj = 0):
    med = stats_dct["med"]
    left = x - horizontal_line_width / 2
    top = stats_dct["whishi"]
    right = x + horizontal_line_width / 2
    bottom = stats_dct["whislo"]
    if top < 1 and top > 0.65:
        manual_adj = 0.23 #-0.05
    elif top < 1:
        manual_adj = -0.05
    else:
        manual_adj = 0
    plt.plot([x, x], [bottom, top], color=color)
    plt.plot([left, right], [top, top], color=color)
    plt.plot([left, right], [bottom, bottom], color=color)
    plt.plot(x, med, 's', color=color_box)
    text = "{:.2f}\n({:.2f}, {:.2f})".format(med, bottom, top)
    plt.text(x, top + 0.27 + manual_adj, text, va = "center", ha = "center", fontsize = 6)

def draw_bxp(stats, labels, headers, fname, color='#2187bb', horizontal_line_width=0.25):
    assert len(stats) == len(labels) and len(labels) == len(headers)
    n = len(labels)
#    fig, ax = plt.subplots()
#    ax.bxp(stats, showmeans=True)
    plt.figure(figsize = (10, 5))
    ylim_max = np.max([max(stats[i][0]["whishi"], stats[i][1]["whishi"]) for i in range(n)]) + 0.8
    for i in range(n):
        top = max(stats[i][0]["whishi"], stats[i][1]["whishi"])
        if top < 1 and top > 0.65:
            manual_adj = 0.2 #-0.05
        elif top < 1:
            manual_adj = -0.05
        else:
            manual_adj = 0
        plot_confidence_interval(i + 1 - 0.25, stats[i][0], color = color, horizontal_line_width = horizontal_line_width, color_box = "#aba9a6", manual_adj = manual_adj)
        plot_confidence_interval(i + 1 + 0.25, stats[i][1], color = color, horizontal_line_width = horizontal_line_width, color_box = "#000000", manual_adj = manual_adj)
        plt.text(i + 1, top + 0.6 + manual_adj, headers[i], va = "center", ha = "center", fontsize = 6, fontweight = "bold")
        if i < n - 1:
            plt.plot([i + 1 + 0.5, i + 1 + 0.5], [0, ylim_max], color = color, linestyle = "--")
#        left = i + 1 - horizontal_line_width / 2
#        right = i + 1 + horizontal_line_width / 2
#        top = stats[i]["whishi"]
#        bottom = stats[i]["whislo"]
#        plt.plot([i + 1, i + 1], [top, bottom], color=color)
#        plt.plot([left, right], [top, top], color=color)
#        plt.plot([left, right], [bottom, bottom], color=color)
#        plt.plot(i + 1, stats[i]["med"], 'o', color='#f44336')
#        plt.text(i + 1, top + 0.2, headers[i], va = "center", ha = "left")
    plt.axhline(y = 1, color = "red", linewidth = 1)
    plt.xticks(list(range(1, n + 1)), [varname_map2(x) for x in labels], fontweight = "bold", fontsize = 8)
    plt.ylabel("Odds Ratio (95% Confidence Interval)", fontweight = "bold", fontsize = 8)
    plt.yticks(np.arange(0, ylim_max, 0.5))
    plt.ylim(bottom = 0, top = ylim_max)
    #plt.title("Figure 1. Associations of Rest-Activity Rhythms with Cardiovascular Disease", loc = "left", fontsize=8, fontweight = "bold")
    plt.grid(False)
    plt.savefig(fname, dpi = 150)
    plt.clf()
    plt.close()

#with open(out_fname, "w") as f:
#    f.write("----- Descriptive Characteristics -----\n")
with open("descriptive_results_2.csv", "r") as f2:
    lines = f2.readlines()
lines = [x.replace('"', "") for x in lines]
cutoff = round(float(lines[1].split(",")[8].strip()), digits)
sample_size = lines[1].split(",")[9].strip()
subhead = None #"\t\tRA Cutoff = " + str(cutoff) + " (Sample Size = " + sample_size + ")"
headers = lines[0].strip().split(",")
headers[1:8] = [headers[1] + f"\n (n={int(sample_size)})"] + [x + f"\n (n={int(int(sample_size) / 2)})" for x in headers[2:8]]
rows = [headers[:9]]
prev_var = ""
tableId = 1
for i in range(1, len(lines)):
    line = lines[i]
    arr = line.strip().split(",")
    if "-" in str(arr[0]):
        curr_var = str(arr[0]).split("-")[0]
        curr_input = curr_var
    else:
        curr_var = ""
        curr_input = str(arr[0])
    if curr_input in output_name_dct_raw:
        input_readable = output_name_dct_raw[curr_input]["Name"]
    else:
        input_readable = varname_map(curr_input)
    if curr_var != "" and curr_var == prev_var:
        for j in range(4):
            if j == 0:
                rows[-1][j] += "\n" + str(arr[j]).replace(curr_input, input_readable)
            else:
                rows[-1][j] += "\n" + str(arr[j])
    else:
        row = [str(arr[0]).replace(curr_input, input_readable), str(arr[1]), str(arr[2]), str(arr[3]), str(arr[4]), str(arr[5]), str(arr[6]), tidy_pval2(arr[7])]
        rows.append(row)
    prev_var = curr_var
headline = f"Table {tableId}. Descriptive Characteristics of the Overall Study Population and by Relative Amplitude Categories"
add_table(word_document, [rows], headline, tableId, subhead = subhead)
footnotes = f"""
1. BMI: body mass index; L5: Least Active 5 hour period; M10: most active 10-hour period; RA: relative amplitude.\n2. RA median value was {cutoff}.\n3. Education Dichotomous: With or without a college degree.\n4. Smoking Status: Smoke or not smoke at all.\n5. Sleep Disorder: Ever diagnosed with sleep disorder or not.\n6. Chi-square tests were used to compare mean values across categories for categorical values and t-tests were used for continuous variables.
"""
p = word_document.add_paragraph()
runner = p.add_run(footnotes.strip())
runner.bold = True
tableId += 1

input_name = "24-h Rest-Activity Pattern Variables"

for reg_model in reg_model_lst:
    for category in output_name_dct[reg_model]:
        df_sub = df[df["RegModel"] == reg_model]
        output_lst = ["self_rated_1", "self_rated_2"]#[x[0] for x in output_name_dct[reg_model][category]] #list(set(df_sub["Output"]))
        input_lst = ["IS60", "IV60", "M10_midpoint", "M10_cnt_100", "L5_midpoint", "L5_cnt", "RA"] #list(set(df_sub["Input"]))
        adj_model_lst = sorted(list(set(df_sub["AdjModel"])))
        rows_lst = []
        bxp_info = {}
        for adj_model in adj_model_lst:
            bxp_info[adj_model] = {}
            
        for output in output_lst:
            for adj_model in adj_model_lst:
                bxp_info[adj_model][output] = {}
                bxp_info[adj_model][output]["stats"] = []
                bxp_info[adj_model][output]["labels"] = []
                bxp_info[adj_model][output]["headers"] = []

            sample_size = df[df["Output"] == output]["SampleSize"].iloc[0]
#            f.write(output + " (Sample Size = " + str(sample_size) + ")\n")
            size_text = " (Sample Size = " + str(sample_size) + ")"
            if reg_model == "linear":
                rows = [[input_name] + ["Model " + str(int(x)) + "\n Est (p-value)\n(95% C.I.)" for x in adj_model_lst]]
                for input in input_lst:
                    row = [varname_map(input)]
                    for adj_model in adj_model_lst:
                        curr = df_sub[(df_sub["Input"] == input) & (df_sub["AdjModel"] == adj_model) & (df_sub["Output"] == output)]
                        if curr.shape[0] > 0:
                            est = curr.iloc[0]["Estimates"]
                            stderr = curr.iloc[0]["Std.Err"]
                            pval = tidy_pval(curr.iloc[0]["P.Value"])
                            content = str(round(est, digits)) + " (" + pval + ")\n" + "(" + str(round(est - ci_coef * stderr, digits)) + " to " + str(round(est + ci_coef * stderr, digits)) + ")"
                        else:
                            content = ""
                        row.append(content)
                    rows.append(row)
            else:
                rows = [[input_name] + ["Model " + str(int(x)) + "\n Tertile1 OR (95% C.I.)\n Tertile2 OR (95% C.I.)\n Tertile3 OR (95% C.I.)\np-trend" for x in adj_model_lst]]
                input_lst = [x + "_tertile" if not x.endswith("_tertile") else x for x in input_lst]
                for input in input_lst:
                    cutoff1 = round(df_cutoff[df_cutoff["Input"] == input.replace("_tertile", "")]["Cutoff1"].iloc[0], digits)
                    cutoff2 = round(df_cutoff[df_cutoff["Input"] == input.replace("_tertile", "")]["Cutoff2"].iloc[0], digits)
                    row = [varname_map(input.replace("_tertile", "")) + f"\n< {cutoff1}\n{cutoff1} - {cutoff2}\n> {cutoff2}"]
                    for adj_model in adj_model_lst:
                        curr = df_sub[(df_sub["Input"] == input) & (df_sub["AdjModel"] == adj_model) & (df_sub["Output"] == output)]
                        if curr.shape[0] > 0:
                            est = curr.iloc[0]["Estimates"]
                            stderr = curr.iloc[0]["Std.Err"]
                            ptrend = tidy_pval(curr.iloc[0]["P.Value"])
#                            or1 = curr.iloc[0]["EstimatesTertile1"]
#                            std1 = curr.iloc[0]["Std.ErrTertile1"]
                            or2 = curr.iloc[0]["EstimatesTertile2"]
                            std2 = curr.iloc[0]["Std.ErrTertile2"]
                            or3 = curr.iloc[0]["EstimatesTertile3"]
                            std3 = curr.iloc[0]["Std.ErrTertile3"]
                            content = str(round(np.exp(or2), digits)) + " (" + str(round(np.exp(or2 - ci_coef * std2), digits)) + "-" + str(round(np.exp(or2 + ci_coef * std2), digits)) + ")\n" + str(round(np.exp(or3), digits)) + " (" + str(round(np.exp(or3 - ci_coef * std3), digits)) + "-" + str(round(np.exp(or3 + ci_coef * std3), digits)) + ")\n" + "" + ptrend
                            
                            stats_dct2 = compute_bxp_stat(or2, std2)
                            stats_dct3 = compute_bxp_stat(or3, std3)
                            bxp_info[adj_model][output]["stats"].append([stats_dct2, stats_dct3])
                            bxp_info[adj_model][output]["labels"].append(input.replace("_tertile", ""))
                            bxp_info[adj_model][output]["headers"].append(ptrend)
                        else:
                            content = ""
                        row.append(content)
                    rows.append(row)
                    if tableId == 2 and output == "cvd_status" and reg_model == "logist":
                        fname = f"Plots/table{tableId}_{output}_model{2}.png"
                        draw_bxp(bxp_info[2][output]["stats"], bxp_info[2][output]["labels"], bxp_info[2][output]["headers"], fname, color= "#cdcac6", horizontal_line_width=0.1)
            if reg_model == "linear":
                model = "Linear"
            else:
                model = "Logistic"
            rows_lst.append(rows)
        headline = f"Table {tableId}. {model} regression models for 24-h Rest-Activity Pattern Variables in Relation to {category}"
        word_document.add_page_break()
        add_table(word_document, rows_lst, headline, tableId, [output_name_dct_raw[x]["Name"] + size_text for x in output_lst])
        if tableId == 2:
            footnotes = """
            1. CI: confidence interval; L5: Least Active 5 hour period; M10: most active 10-hour period; OR: Odds Ratio.\n2. Model 1: Represents univariate associations.\n3. Model 2: Adjusted for age, sex, marital status, education, race/ethnicity, smoking, and alcohol use.\n4. Model 3: Adjusted for model 2 covariates and additionally for BMI.\n5. Model 4: Adjusted for model 2 covariates and additionally for self-reported sleep disorders.
            """
            p = word_document.add_paragraph()
            runner = p.add_run(footnotes.strip())
            runner.bold = True
        tableId += 1

word_document.save(out_fname)
